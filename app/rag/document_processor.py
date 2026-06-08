"""
Enterprise Agentic RAG Assistant
Document processor — extracts text from PDF, DOCX, and TXT files
while preserving per-page metadata and splitting into intelligent chunks.
"""

from __future__ import annotations

import hashlib
import io
from pathlib import Path
from typing import Iterator

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

from app.models.schemas import ChunkMetadata, FileType
from app.utils.config import get_settings
from app.utils.logger import get_logger, log_latency

logger = get_logger(__name__)
settings = get_settings()


# ── Helper: file-type detection ───────────────────────────────────────────────

def _detect_file_type(path: Path) -> FileType:
    """
    Infer FileType from the file extension.

    Args:
        path: Filesystem path to the uploaded file.

    Returns:
        FileType enum member.

    Raises:
        ValueError: If the extension is not supported.
    """
    ext = path.suffix.lower().lstrip(".")
    mapping = {"pdf": FileType.PDF, "docx": FileType.DOCX, "txt": FileType.TXT}
    if ext not in mapping:
        raise ValueError(
            f"Unsupported file type '.{ext}'. Supported: {list(mapping)}"
        )
    return mapping[ext]


# ── Per-format extractors ─────────────────────────────────────────────────────

def _extract_pdf(path: Path) -> list[tuple[int, str]]:
    """
    Extract (page_number, text) tuples from a PDF.
    Uses pypdf for pure-Python parsing without any native dependencies.

    Args:
        path: Path to the PDF file.

    Returns:
        List of (1-indexed page number, page text) tuples.
    """
    try:
        import pypdf  # lazy import — not needed for non-PDF files
    except ImportError as exc:
        raise ImportError(
            "pypdf is required for PDF processing. "
            "Install it with: pip install pypdf"
        ) from exc

    pages: list[tuple[int, str]] = []
    with open(path, "rb") as fh:
        reader = pypdf.PdfReader(fh)
        for idx, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            pages.append((idx, text))

    logger.debug(
        "PDF extracted",
        extra={"file": path.name, "num_pages": len(pages)},
    )
    return pages


def _extract_docx(path: Path) -> list[tuple[int, str]]:
    """
    Extract text from a DOCX file.
    DOCX files do not have a native concept of pages; we group paragraphs
    into virtual pages of roughly 3 000 characters to preserve a meaningful
    page number in metadata.

    Args:
        path: Path to the DOCX file.

    Returns:
        List of (virtual_page_number, text) tuples.
    """
    try:
        import docx  # lazy import
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for DOCX processing. "
            "Install it with: pip install python-docx"
        ) from exc

    PAGE_SIZE = 3_000  # characters per virtual page
    doc = docx.Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)

    pages: list[tuple[int, str]] = []
    for page_num, start in enumerate(range(0, max(len(full_text), 1), PAGE_SIZE), start=1):
        chunk = full_text[start : start + PAGE_SIZE]
        if chunk:
            pages.append((page_num, chunk))

    logger.debug(
        "DOCX extracted",
        extra={"file": path.name, "virtual_pages": len(pages)},
    )
    return pages


def _extract_txt(path: Path) -> list[tuple[int, str]]:
    """
    Extract text from a plain-text file.
    The entire file is returned as page 1 (no page concept for TXT).

    Args:
        path: Path to the TXT file.

    Returns:
        Single-element list [(1, full_text)].
    """
    text = path.read_text(encoding="utf-8", errors="replace")
    logger.debug("TXT extracted", extra={"file": path.name, "chars": len(text)})
    return [(1, text)]


# ── Chunking helpers ──────────────────────────────────────────────────────────

def _build_splitter() -> RecursiveCharacterTextSplitter:
    """
    Build a RecursiveCharacterTextSplitter configured from application settings.
    Splits on paragraph/sentence/word/character boundaries in that order.
    """
    return RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
        length_function=len,
        is_separator_regex=False,
    )


def _chunk_id(source: str, page: int, index: int) -> str:
    """Generate a deterministic chunk ID from source, page, and index."""
    raw = f"{source}::{page}::{index}"
    return "chunk_" + hashlib.md5(raw.encode()).hexdigest()[:12]


# ── Public API ────────────────────────────────────────────────────────────────

class DocumentProcessor:
    """
    Orchestrates document loading, text extraction, and intelligent chunking.

    Usage::

        processor = DocumentProcessor()
        chunks, metadata = processor.process(path="report.pdf", document_id="abc")
        # chunks is a list[Document] ready for embedding & vector storage
    """

    def __init__(self) -> None:
        self._splitter = _build_splitter()

    def process(
        self,
        path: str | Path,
        document_id: str,
    ) -> tuple[list[Document], list[ChunkMetadata]]:
        """
        Load a document, extract text, and split into chunks.

        Args:
            path:        Filesystem path to the uploaded document.
            document_id: UUID assigned to this document in the registry.

        Returns:
            Tuple of:
            - ``list[Document]``: LangChain Document objects with text + metadata.
            - ``list[ChunkMetadata]``: Pydantic metadata objects matching each Document.

        Raises:
            ValueError: For unsupported file types.
            FileNotFoundError: If the file does not exist.
        """
        file_path = Path(path)
        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")

        file_type = _detect_file_type(file_path)

        with log_latency(
            logger,
            "document_processing",
            file_name=file_path.name,
            file_type=file_type.value,
        ):
            pages = self._extract_pages(file_path, file_type)
            documents, chunk_metas = self._chunk_pages(
                pages, file_path.name, document_id
            )

        logger.info(
            "Document processed",
            extra={
                "file_name": file_path.name,
                "document_id": document_id,
                "num_pages": len(pages),
                "num_chunks": len(documents),
            },
        )
        return documents, chunk_metas

    # ── Internal helpers ──────────────────────────────────────────────────────

    def _extract_pages(
        self, path: Path, file_type: FileType
    ) -> list[tuple[int, str]]:
        """Dispatch to the appropriate per-format extractor."""
        extractors = {
            FileType.PDF: _extract_pdf,
            FileType.DOCX: _extract_docx,
            FileType.TXT: _extract_txt,
        }
        return extractors[file_type](path)

    def _chunk_pages(
        self,
        pages: list[tuple[int, str]],
        filename: str,
        document_id: str,
    ) -> tuple[list[Document], list[ChunkMetadata]]:
        """
        Split each page's text using the configured splitter, assign metadata,
        and return parallel lists of Document and ChunkMetadata objects.
        """
        all_docs: list[Document] = []
        all_metas: list[ChunkMetadata] = []
        global_index = 0
        total_chunks_estimate = sum(
            max(1, len(text) // settings.chunk_size) for _, text in pages
        )

        for page_num, page_text in pages:
            if not page_text.strip():
                continue  # Skip empty pages

            sub_chunks = self._splitter.split_text(page_text)
            for sub_idx, chunk_text in enumerate(sub_chunks):
                if not chunk_text.strip():
                    continue

                cid = _chunk_id(filename, page_num, global_index)
                meta = ChunkMetadata(
                    source=filename,
                    page=page_num,
                    chunk_id=cid,
                    document_id=document_id,
                    total_chunks=total_chunks_estimate,
                )
                doc = Document(
                    page_content=chunk_text,
                    metadata=meta.model_dump(),
                )
                all_docs.append(doc)
                all_metas.append(meta)
                global_index += 1

        return all_docs, all_metas

    def iter_chunks(
        self,
        path: str | Path,
        document_id: str,
    ) -> Iterator[tuple[Document, ChunkMetadata]]:
        """
        Generator version of :meth:`process` — yields one (Document, metadata)
        pair at a time, useful for streaming large documents into the vector store.
        """
        docs, metas = self.process(path, document_id)
        yield from zip(docs, metas)
